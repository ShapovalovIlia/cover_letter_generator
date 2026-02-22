import io
import os

import docx
import pymupdf
import pytest

os.environ.setdefault("OPENAI_API_KEY", "sk-test-fake-key")


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    with pymupdf.open() as doc:  # type: ignore[no-untyped-call]
        page = doc.new_page()
        page.insert_text(
            (72, 72), "John Doe\nSoftware Engineer\n5 years experience"
        )
        return doc.tobytes()  # type: ignore[no-any-return]


@pytest.fixture
def sample_docx_bytes() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("Product Manager")
    doc.add_paragraph("10 years experience")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def empty_pdf_bytes() -> bytes:
    with pymupdf.open() as doc:  # type: ignore[no-untyped-call]
        doc.new_page()
        return doc.tobytes()  # type: ignore[no-any-return]
